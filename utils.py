from docx import Document
from docx.shared import Pt

def collect_user_data():
    print("🔹 Let's start with your personal details:\n")
    
    personal = {
        "name": input("Full Name: "),
        "email": input("Email: "),
        "phone": input("Phone Number: "),
        "linkedin": input("LinkedIn URL: "),
        "github": input("GitHub URL: ")
    }

    print("\n🔹 Let's enter your education details:\n")
    
    education = []
    while True:
        degree = input("Degree / Program: ")
        institution = input("Institution Name: ")
        start = input("Start Year: ")
        end = input("End Year: ")

        education.append({
            "degree": degree,
            "institution": institution,
            "start": start,
            "end": end
        })

        more = input("Do you want to add another education entry? (yes/no): ").strip().lower()
        if more != 'yes':
            break

        print("\n🔹 Let's enter your work experience details:\n")

    experience = []
    while True:
        title = input("Job Title: ")
        company = input("Company Name: ")
        start = input("Start Date (e.g., Jan 2022): ")
        end = input("End Date (e.g., Present or Dec 2023): ")

        print("Enter job description bullet points (press Enter twice to finish):")
        description = []
        while True:
            point = input("- ")
            if point.strip() == "":
                break
            description.append(point)

        experience.append({
            "title": title,
            "company": company,
            "start": start,
            "end": end,
            "description": description
        })

        more = input("Do you want to add another job? (yes/no): ").strip().lower()
        if more != 'yes':
            break
    
    print("\n🔹 Let's enter your project details:\n")

    projects = []
    while True:
        title = input("Project Title: ")
        description = input("Short Description: ")
        tech = input("Technologies Used (comma-separated): ")
        link = input("Project Link (GitHub/Live URL, optional): ")

        projects.append({
            "title": title,
            "description": description,
            "technologies": [t.strip() for t in tech.split(",") if t.strip()],
            "link": link
        })

        more = input("Do you want to add another project? (yes/no): ").strip().lower()
        if more != 'yes':
            break


    print("\n🔹 Let's enter your skills:\n")
    skills_input = input("Enter your skills (comma-separated): ")
    skills = [skill.strip() for skill in skills_input.split(",") if skill.strip()]

    print("\n🔹 Extras: Languages & Certifications\n")
    languages_input = input("Languages Known (comma-separated): ")
    certifications_input = input("Certifications (comma-separated): ")

    extras = {
        "languages": [lang.strip() for lang in languages_input.split(",") if lang.strip()],
        "certifications": [cert.strip() for cert in certifications_input.split(",") if cert.strip()]
    }

        # Return all collected data in a dictionary
    return {
        "personal": personal,
        "education": education,
        "experience": experience,
        "projects": projects,
        "skills": skills,
        "extras": extras
    }

   

def generate_docx(data, filename="outputs/resume.docx"):
    doc = Document()

    # Personal Details
    personal = data["personal"]
    doc.add_heading(personal["name"], 0)
    doc.add_paragraph(f"Email: {personal['email']}")
    doc.add_paragraph(f"Phone: {personal['phone']}")
    doc.add_paragraph(f"LinkedIn: {personal['linkedin']}")
    doc.add_paragraph(f"GitHub: {personal['github']}")

    doc.add_heading("Education", level=1)
    for edu in data["education"]:
        doc.add_paragraph(f"{edu['degree']} - {edu['institution']} ({edu['start']} - {edu['end']})")

    doc.add_heading("Work Experience", level=1)
    for job in data["experience"]:
        doc.add_paragraph(f"{job['title']} - {job['company']} ({job['start']} - {job['end']})")
        for point in job["description"]:
            doc.add_paragraph(f"• {point}", style='List Bullet')

    doc.add_heading("Projects", level=1)
    for proj in data["projects"]:
        doc.add_paragraph(f"{proj['title']}: {proj['description']}")
        doc.add_paragraph(f"Technologies: {', '.join(proj['technologies'])}")
        if proj["link"]:
            doc.add_paragraph(f"Link: {proj['link']}")

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(", ".join(data["skills"]))

    doc.add_heading("Extras", level=1)
    doc.add_paragraph(f"Languages: {', '.join(data['extras']['languages'])}")
    doc.add_paragraph(f"Certifications: {', '.join(data['extras']['certifications'])}")

    doc.save(filename)
    print(f"\n📄 Resume saved as: {filename}")



